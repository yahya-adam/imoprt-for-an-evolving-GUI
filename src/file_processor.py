import io
import os
import sys
import requests
from PyPDF2 import PdfReader
from docx import Document

# Global variable for magic instance
_magic_instance = None

def get_magic_instance():
    """Create and configure a magic instance with explicit database path"""
    global _magic_instance
    if _magic_instance is None:
        import magic
        try:
            if getattr(sys, 'frozen', False):
                # Running in PyInstaller bundle
                base_path = sys._MEIPASS
                magic_db_path = os.path.join(base_path, "magic.mgc")
                
                if os.path.exists(magic_db_path):
                    print(f"Using magic database at: {magic_db_path}")
                    # Create instance that returns MIME types
                    _magic_instance = magic.Magic(mime=True, magic_file=magic_db_path)
                else:
                    print("WARNING: magic.mgc not found in bundle. Using default.")
                    _magic_instance = magic.Magic(mime=True)
            else:
                # Running in development mode
                magic_db_path = os.path.join(os.path.dirname(magic.__file__), "libmagic", "magic.mgc")
                if os.path.exists(magic_db_path):
                    print(f"Using magic database at: {magic_db_path}")
                    _magic_instance = magic.Magic(mime=True)
                    _magic_instance.file = magic_db_path
                else:
                    print("WARNING: magic.mgc not found in development. Using default.")
                    _magic_instance = magic.Magic(mime=True)
        except Exception as e:
            print(f"ERROR: Failed to create magic instance: {str(e)}")
            raise
    return _magic_instance

def extract_text(file_bytes, filename):
    try:
        m = get_magic_instance()
        # Get MIME type instead of description
        file_type = m.from_buffer(file_bytes)
        
        # Normalize MIME types for consistent comparison
        normalized_file_type = file_type.lower()
        
        if "pdf" in normalized_file_type:
            return extract_text_from_pdf(file_bytes)
        elif "word" in normalized_file_type or \
             "vnd.openxmlformats-officedocument.wordprocessingml.document" in normalized_file_type:
            return extract_text_from_docx(file_bytes)
        else:
            raise ValueError(f"Unsupported file type: {file_type} for {filename}")
            
    except Exception as e:
        raise ValueError(f"Text extraction failed for {filename}: {str(e)}")
    
def extract_text_from_pdf(file_bytes):
    reader = PdfReader(io.BytesIO(file_bytes))
    text = ""
    for page in reader.pages:
        text += page.extract_text() + "\n"
    return text

def extract_text_from_docx(file_bytes):
    doc = Document(io.BytesIO(file_bytes))
    return "\n".join([para.text for para in doc.paragraphs])

def get_file_metadata(file_bytes, filename):
    m = get_magic_instance()
    # Get MIME type instead of description
    file_type = m.from_buffer(file_bytes)
    size = len(file_bytes)
    metadata = {"size": size, "filename": filename}
    
    try:
        # Normalize MIME types for consistent comparison
        normalized_file_type = file_type.lower()
        
        if "pdf" in normalized_file_type:
            reader = PdfReader(io.BytesIO(file_bytes))
            metadata.update({
                "author": reader.metadata.get("/Author", "Unknown"),
                "title": reader.metadata.get("/Title", filename),
                "created": reader.metadata.get("/CreationDate", ""),
                "modified": reader.metadata.get("/ModDate", ""),
                "pages": len(reader.pages)
            })
        elif "word" in normalized_file_type or \
             "vnd.openxmlformats-officedocument.wordprocessingml.document" in normalized_file_type:
            doc = Document(io.BytesIO(file_bytes))
            metadata.update({
                "author": doc.core_properties.author or "Unknown",
                "created": str(doc.core_properties.created) if doc.core_properties.created else "",
                "modified": str(doc.core_properties.modified) if doc.core_properties.modified else "",
                "revision": doc.core_properties.revision
            })
    except Exception:
        # Metadata extraction failed, but we still have basic info
        pass
    
    return metadata

def validate_file(file_bytes, filename):
    """Validate file type and size"""
    MAX_SIZE = 10 * 1024 * 1024  # 10MB
    
    if len(file_bytes) > MAX_SIZE:
        raise ValueError(f"File too large (max {MAX_SIZE//1024//1024}MB)")
    
    m = get_magic_instance()
    # Get MIME type instead of description
    file_type = m.from_buffer(file_bytes)
    
    # Check for PDF or Word documents by MIME type patterns
    normalized_file_type = file_type.lower()
    allowed_patterns = [
        "pdf",
        "msword",
        "vnd.openxmlformats-officedocument.wordprocessingml.document"
    ]
    
    if not any(pattern in normalized_file_type for pattern in allowed_patterns):
        raise ValueError(f"Unsupported file type: {file_type}")
    
    return True



def scan_for_viruses(file_bytes):
    """Scan file using VirusTotal API (requires API key)"""
    API_KEY = "YOUR_VIRUSTOTAL_API_KEY"
    url = "https://www.virustotal.com/api/v3/files"
    
    headers = {
        "x-apikey": API_KEY
    }
    
    files = {"file": ("document", file_bytes)}
    response = requests.post(url, headers=headers, files=files)
    
    if response.status_code == 200:
        result = response.json()
        if result["data"]["attributes"]["last_analysis_stats"]["malicious"] > 0:
            raise ValueError("File contains malware!")
    else:
        # Fallback: log but don't block
        print("VirusTotal scan failed, proceeding anyway")